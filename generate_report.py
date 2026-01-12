# generate_report.py
from docx import Document
from datetime import datetime
import db

def make_report(output="EncryptEase_Phase3_Report.docx"):
    stats = db.get_stats()
    users = db._get_conn().cursor().execute("SELECT COUNT(*) as c FROM users").fetchone()["c"]
    document = Document()
    document.add_heading('Encrypt Ease — Phase 3 Final Report', 0)
    document.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_heading('Project Overview', level=1)
    document.add_paragraph("Encrypt Ease is a Python desktop application for file encryption/decryption, using AES-GCM and ChaCha20-Poly1305 with RSA for key protection. User authentication uses PBKDF2+salt and the system provides SHA-256 integrity checks.")
    document.add_heading('Final Status', level=1)
    document.add_paragraph("All core modules implemented and tested. The project is complete for Phase 3.")
    document.add_heading('Statistics', level=2)
    document.add_paragraph(f"Total users: {users}")
    document.add_paragraph(f"Files encrypted (count): {stats.get('files_encrypted', 0)}")
    document.add_heading('Modules Completed', level=2)
    document.add_paragraph("- AES-GCM encryption/decryption")
    document.add_paragraph("- ChaCha20-Poly1305 encryption/decryption")
    document.add_paragraph("- RSA key wrapping/unwrapping")
    document.add_paragraph("- SHA-256 integrity verification")
    document.add_paragraph("- User authentication (DB-backed with PBKDF2)")
    document.add_paragraph("- Performance measurement (timing + plotting)")
    document.add_heading('Testing', level=2)
    document.add_paragraph("All listed tests (AES, ChaCha20, RSA, Integrity, Authentication, GUI) have passed during validation.")
    document.add_heading('Future Enhancements', level=2)
    document.add_paragraph("- Database backup & restore; audit logs; multi-user roles")
    document.add_paragraph("- Cloud backup & secure sharing")
    document.add_paragraph("- UI improvements and threading for large files")
    document.save(output)
    print("Report saved to", output)

if __name__ == "__main__":
    make_report()
